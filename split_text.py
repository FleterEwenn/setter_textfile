from docx import Document
from odf.opendocument import load
from odf.text import P

class File_text:
	def __init__(self, name:str):
		self.file_name = name

		self.isdocx = False
		self.isodt = False
		self.document = None

		if self.file_name.lower().endswith('.docx'):
			self.isdocx = True
		elif self.file_name.lower().endswith('.odt'):
			self.isodt = True

		self.text = self.read_text()
		self.current_words = self.split_paragraphs()
	
	def read_text(self)->list:
		text = []

		if self.isdocx:
			self.document = Document(self.file_name)
			for para in self.document.paragraphs:
				text.append(para.text)

		elif self.isodt:
			self.document = load(self.file_name)
			for para in self.document.getElementsByType(P):
				text.append(para.firstChild.data)
		
		else:
			with open(self.file_name, 'r') as file:
				text.append(file.read())

		return text
	
	def split_words(self, text:str)->list:
		"""
		Split properly text with space and also split  with punctuation
		"""
		text_toreturn = []
		chars = [',', '.', ';']
		for str_ in text.split(' '):
			try:
				if str_[-1] in chars:
					text_toreturn.append(str_[:-1])
					text_toreturn.append(str_[-1])
				else:
					text_toreturn.append(str_)

			except:
				text_toreturn.append(str_)
		return text_toreturn
	
	def split_paragraphs(self)->list:
		words_list = []

		if self.isdocx or self.isodt:
			for para in self.text:
				words_list.append(self.split_words(para[0]))
		
		else:
			words_list.append(self.split_words(self.text[0]))
		
		return words_list
	
	def save(self, output):
		if self.isdocx:

			for i in range(len(self.document.paragraphs)):
				self.document.paragraphs[i].runs.clear()

			for para in self.current_words:
				text_tosave = ''
				for elt in para:
					if elt in [',', ';', '.']:
						text_tosave += elt
					else:
						text_tosave += ' ' + elt
				self.document.add_paragraph(text_tosave)

			self.document.save(output)

		elif self.isodt:

			for child in list(self.document.text.childNodes):
				self.document.text.removeChild(child)

			for para in self.current_words:
				text_tosave = ''
				for elt in para:
					if elt in [',', ';', '.']:
						text_tosave += elt
					else:
						text_tosave += ' ' + elt

				para_tosave = P(text=text_tosave)
				self.document.text.addElement(para_tosave)
			
			self.document.save(output)
		
		else:
			with open(output, 'w') as file:
				for elt in self.current_words[0]:
					if elt in [',', ';', '.']:
						file.write(elt)
					else:
						file.write(' ' + elt)