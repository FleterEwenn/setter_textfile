import argparse
import os
import string
from docx import Document
from odf.opendocument import load
from odf.text import P
from pipeline_commands import Upper, Lower, Replace
from split_text import Split_text

parser = argparse.ArgumentParser(prog="textpipe", description="Moteur de traitement de texte programmable ")

parser.add_argument('-o', '--output', help='fichier de sortie')
parser.add_argument('-i', '--input', help="fichier d'entrée")
parser.add_argument('-p', '--pipeline', help='pipeline de commande replace:old:new/lower:word/upper:word/stats/capitalize:word ...')

args = parser.parse_args()

if not args.input:
	raise Exception("il manque un fichier d'entrée")

elif not os.path.isfile(args.input):
	raise Exception(f"le fichier d'entrée {args.input} mis en paramètre n'existe pas")

if str(args.input).lower().endswith('.docx'):
	document = Document(args.input)
	current_words = []
	for para in document.paragraphs:
		current_words.append(Split_text(para.text))

elif str(args.input).lower().endswith('.txt'):
	with open(args.input, 'r') as file:
		text = file.read()
	current_words = Split_text(text)

elif str(args.input).lower().endswith('.odt'):
	document_odt = load(args.input)
	current_words = []
	for para in document_odt.getElementsByType(P):
		current_words.append(Split_text(para.firstChild.data))

else:
	raise Exception(f"l'extension du fichier d'entrée {args.output} n'est pas traitable")

if args.pipeline:

	cmds = [cmd for cmd in args.pipeline.split('/')]

	for cmd in cmds:

		if cmd.startswith('stats'):
			len_word = 0
			len_punctuation = 0
			if args.input.lower().endswith('.docx') or args.input.lower().endswith('.odt'):
				for para in current_words:
					for elt in para:
						if elt.isalpha():
							len_word += 1
						elif elt in string.punctuation:
							len_punctuation += 1
			else:
				for elt in current_words:
					if elt.isalpha():
						len_word += 1
					elif elt in string.punctuation:
						len_punctuation += 1
			
			print("words : " + str(len_word))
			print("punctuation : " + str(len_punctuation))
		
		elif cmd.startswith('wordcount'):
			try:
				current_cmd, word_to_count = cmd.split(':')
			except:
				raise Exception("la commande {cmd} ne correspond pas au schéma wordcount:word")
			
			word_iteration = 0
			if args.input.lower().endswith('.docx') or args.input.lower().endswith('.odt'):
				for para in current_words:
					for elt in para:
						if elt == word_to_count:
							word_iteration += 1
			else:
				for elt in current_words:
					if elt == word_to_count:
						word_iteration += 1
			
			print('word ' + '"' + word_to_count+ '"' + ' : ' + str(word_iteration))

		elif cmd.startswith('replace'):

			try:
				current_cmd, old, new = cmd.split(':')
			except:
				raise Exception(f"la commande {cmd} ne correspond pas au schéma replace:old:new")

			if args.input.lower().endswith('.docx') or args.input.lower().endswith('.odt'):
				for i in range(len(current_words)):
					current_words[i] = Replace(old, new, current_words[i])

			else:
				current_words = Replace(old, new, current_words)

		elif cmd.startswith('upper'):

			if len(cmd.split(':')) == 2:
				current_cmd, word_to_upper = cmd.split(':')
				
				if args.input.lower().endswith('.docx') or args.input.lower().endswith('.odt'):
					for i in range(len(current_words)):
						current_words[i] = Upper(current_words[i], word_to_upper)
				else:
					current_words = Upper(current_words, word_to_upper)

			elif len(cmd.split(':')) == 3:
				current_cmd, word_target, letter_to_upper = cmd.split(':')

				if args.input.lower().endswith('.docx'):
					for i in range(len(current_words)):
						current_words[i] = Upper(current_words[i], word_target, letter=letter_to_upper)
				else:
					current_words = Upper(current_words, word_target, letter=letter_to_upper)

		elif cmd.startswith('lower'):

			if len(cmd.spit(':')) == 2:
				current_cmd, word_to_lower = cmd.split(':')
				if args.input.lower().endswith('.docx') or args.input.lower().endswith('.odt'):
					for i in range(len(current_words)):
						current_words[i] = Lower(current_words[i], word_to_lower)
				else:
					current_words = Lower(current_words, word_to_lower)

			elif len(cmd.split(':')) == 3:
				current_cmd, word_target, letter_to_upper = cmd.split(':')
				if args.input.lower().endswith('.docx') or args.input.lower().endswith('.odt'):
					for i in range(len(current_words)):
						current_words[i] = Lower(current_words[i], word_target, letter=letter_to_upper)
				else:
					current_words = Lower(current_words, word_target, letter=letter_to_upper)
		
		elif cmd.startswith('capitalize'):

			try:
				current_cmd, word_to_capitalize = cmd.split(':')
			except:
				raise Exception(f"la commande {cmd} ne correspond pas au schéma capitalize:word")

			if args.input.lower().endswith('.docx'):
				for i in range(len(current_words)):
					current_words[i] = Upper(current_words[i], word_to_capitalize, letter=word_to_capitalize[0])
			else:
				current_words = Upper(current_words, word_to_capitalize, letter=word_to_capitalize[0])

		else:
			raise Exception(f"la commande {cmd} n'exite pas")

if args.input.lower().endswith('.docx'):
	for i in range(len(document.paragraphs)):
		document.paragraphs[i].text = ''
		document.paragraphs[i].runs.clear()
		for elt in current_words[i]:
			if elt in [',', ';', '.']:
				document.paragraphs[i].text += elt
			else:
				document.paragraphs[i].text += ' ' + elt

	if args.output:
		document.save(args.output)
	else:
		document.save('output.docx')

elif args.input.lower().endswith('.odt'):

	document_odt = load(args.input)

	for child in list(document_odt.text.childNodes):
		document_odt.text.removeChild(child)

	for para in current_words:
		text_tosave = ''
		for elt in para:
			if elt in [',', ';', '.']:
				text_tosave += elt
			else:
				text_tosave += ' ' + elt
		para_tosave = P(text=text_tosave)
		document_odt.text.addElement(para_tosave)

	if args.output:
		document_odt.save(args.output)
	else:
		document_odt.save('output.odt')

else:
	if args.output:
		with open(args.output, 'w', encoding='utf-8') as file:
			for elt in current_words:
				if elt in [',', ';', '.']:
					file.write(elt)
				else:
					file.write(' ' + elt)

	else:
		with open('output.txt', 'w', encoding='utf-8') as file:
			for elt in current_words:
				if elt in [',', ';', '.']:
					file.write(elt)
				else:
					file.write(' ' + elt)