import argparse
import os
import string
from docx import Document

parser = argparse.ArgumentParser(prog="textpipe", description="Moteur de traitement de texte programmable ")

parser.add_argument('-o', '--output', help='fichier de sortie')
parser.add_argument('-i', '--input', help="fichier d'entrée")
parser.add_argument('-p', '--pipeline', help='pipeline de commande replace:old:new/lower:word/upper:word/stats ...')

args = parser.parse_args()

def Replace(old:str, new:str, words:list)->list:
	"""
	Change the old string to the new in the variable words
	"""
	
	for i in range(len(words)):
		if words[i] == old:
			words[i] = new
	return words	
	
			
def Split_text(text:str)->str:
	"""
	Split properly text with space and also split punctuation
	"""
	text_toreturn = []
	chars = [',', '.', ';']
	for str_ in text.split(' '):
		if str_[-1] in chars:
			text_toreturn.append(str_[:-1])
			text_toreturn.append(str_[-1])

		else:
			text_toreturn.append(str_)
	return text_toreturn

def Upper(words:list, wordtoupper:str, letter=None)->list:
	"""
	upper the word 'wordtoupper'
	"""

	for i in range(len(words)):
		if words[i] == wordtoupper:
			if letter:
				for j in range(len(words[i])):
					if words[i][j] == letter:
						words[i] = words[i][:j] + words[i][j].upper() + words[i][j+1:]
			else:
				words[i] = words[i].upper()

	return words

def Lower(words:list, wordtolower:str, letter=None)->list:
	"""
	lower the word 'wordtolower'
	"""
	for i in range(len(words)):
		if words[i].lower() == wordtolower:
			if letter:
				for j in range(len(words[i])):
					if words[i][j] == letter:
						words[i] = words[i][:j] + words[i][j].lower() + words[i][j+1:]
			else:
				words[i] = words[i].lower()
	return words			

if not args.input:
	raise Exception("il manque un fichier d'entrée ou de sortie")

elif not os.path.isfile(args.input):
	raise Exception("les fichiers mis en paramètre n'existent pas")

if str(args.input).lower().endswith('.docx'):
	document = Document(args.input)
	current_words = []
	for para in document.paragraphs:
		current_words.append(Split_text(para.text))

elif str(args.input).lower().endswith('.txt'):
	with open(args.input, 'r') as file:
		text = file.read()

	current_words = Split_text(text)
else:
	raise Exception("l'extension du fichier d'entrée n'est pas traitable")

if args.pipeline:

	cmds = [cmd for cmd in args.pipeline.split('/')]

	for cmd in cmds:

		if cmd.startswith('stats'):
			len_word = 0
			len_punctuation = 0
			if args.input.lower().endswith('.docx'):
				for para in current_words:
					for elt in para:
						if elt.isalpha():
							len_word += 1
						elif elt in string.punctuation:
							len_punctuation += 1
			else:
				for elt in current_words:
					if elt.isalpha():
						len_word += 1
					elif elt in string.punctuation:
						len_punctuation += 1
			
			print("words : " + str(len_word))
			print("punctuation : " + str(len_punctuation))

		if cmd.startswith('replace'):

			current_cmd, old, new = cmd.split(':')

			if args.input.lower().endswith('.docx'):
				for i in range(len(current_words)):
					current_words[i] = Replace(old, new, current_words[i])
					
			else:
				current_words = Replace(old, new, current_words)
		
		if cmd.startswith('upper'):

			if len(cmd.split(':')) == 2:
				current_cmd, word_to_upper = cmd.split(':')
				
				if args.input.lower().endswith('.docx'):
					for i in range(len(current_words)):
						current_words[i] = Upper(current_words[i], word_to_upper)
				else:
					current_words = Upper(current_words, word_to_upper)
			elif len(cmd.split(':')) == 3:
				current_cmd, word_target, letter_to_upper = cmd.split(':')
				
				if args.input.lower().endswith('.docx'):
					for i in range(len(current_words)):
						current_words[i] = Upper(current_words[i], word_target, letter=letter_to_upper)
				else:
					current_words = Upper(current_words, word_target, letter=letter_to_upper)
		
		if cmd.startswith('lower'):
			if len(cmd.spit(':')) == 2:
				current_cmd, word_to_lower = cmd.split(':')
				if args.input.lower().endswith('.docx'):
					for i in range(len(current_words)):
						current_words[i] = Lower(current_words[i], word_to_lower)
				else:
					current_words = Lower(current_words, word_to_lower)
			elif len(cmd.split(':')) == 3:
				current_cmd, word_target, letter_to_upper = cmd.split(':')
				if args.input.lower().endswith('.docx'):
					for i in range(len(current_words)):
						current_words[i] = Lower(current_words[i], word_target, letter=letter_to_upper)
				else:
					current_words = Lower(current_words, word_target, letter=letter_to_upper)

if args.input.lower().endswith('.docx'):
	for i in range(len(document.paragraphs)):
		document.paragraphs[i].text = ' '
		for word in current_words[i]:
			if word in [',', ';', '.']:
				document.paragraphs[i].text += word
			else:
				document.paragraphs[i].text += ' ' + word

	if args.output:
		document.save(args.output)
	else:
		document.save('output.docx')

else:
	if args.output:
		with open(args.output, 'w') as file:
			for word in current_words:
				if word in [',', ';', '.']:
					file.write(word)
				else:
					file.write(' ' + word)

	else:
		with open('output.txt', 'w') as file:
			for word in current_words:
				if word in [',', ';', '.']:
					file.write(word)
				else:
					file.write(' ' + word)